# -*- coding: utf-8 -*-
"""
Gera docs/Atividade2_Quimiochat.docx
Executa com: backend\\venv\\Scripts\\python.exe generate_atividade2.py
"""
import os
import re
import ftfy
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOCS_DIR = os.path.join(BASE_DIR, "docs")

# ─────────────────────────────────────────────────────────────────────────────
# 1. LEITURA E CONSERTO DE ENCODING
#    ftfy.fix_text() detecta e corrige automaticamente double-encoding,
#    mojibake e outros artefatos de charset — é a abordagem mais robusta.
# ─────────────────────────────────────────────────────────────────────────────
def read_fixed(path):
    with open(path, "rb") as f:
        raw = f.read()
    # Tenta utf-8 puro; se falhar, usa latin-1 (lê qualquer byte sem erro)
    try:
        text = raw.decode("utf-8")
    except UnicodeDecodeError:
        text = raw.decode("latin-1")
    # ftfy corrige double-encoding, sequências mojibake etc.
    return ftfy.fix_text(text)


# ─────────────────────────────────────────────────────────────────────────────
# 2. LIMPEZA GERAL DO TEXTO
# ─────────────────────────────────────────────────────────────────────────────
def clean_text(text):
    # Remove badges GitHub: [![label](img)](link)
    text = re.sub(r"\[!\[.*?\]\(.*?\)\]\(.*?\)", "", text)
    # Remove emojis e símbolos não-textuais
    text = re.sub(
        "["
        "\U0001F300-\U0001FAFF"
        "\U00002702-\U000027B0"
        "\U000024C2-\U0001F251"
        "\U0000200D\U0000FE0F"
        "]+",
        "",
        text,
        flags=re.UNICODE,
    )
    # Normaliza quebras de linha Windows → Unix
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Remove linhas que são apenas separadores de markdown (---) sem texto útil
    # mas preserva o conteúdo seguinte — evita o bug "---# Heading"
    text = re.sub(r"^-{3,}\s*$", "", text, flags=re.MULTILINE)
    return text


# ─────────────────────────────────────────────────────────────────────────────
# 3. LEITURA DOS ARQUIVOS
# ─────────────────────────────────────────────────────────────────────────────
MODELO_LOGICO = clean_text(read_fixed(os.path.join(DOCS_DIR, "modelo_logico.md")))
MODELO_FISICO = clean_text(read_fixed(os.path.join(DOCS_DIR, "modelo_fisico.sql")))
API_DOCS      = clean_text(read_fixed(os.path.join(DOCS_DIR, "API_DOCUMENTATION.md")))
README        = clean_text(read_fixed(os.path.join(BASE_DIR, "README.md")))
DER_PATH      = os.path.join(DOCS_DIR, "DER.png")


# ─────────────────────────────────────────────────────────────────────────────
# 4. HELPERS DE FORMATAÇÃO DOCX
# ─────────────────────────────────────────────────────────────────────────────
def add_heading(doc, text, level=1):
    text = text.strip()
    if not text:
        return
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return p


def add_body(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text.strip())
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(12)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    return p


def add_code_block(doc, code_text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F0F0F0")
    pPr.append(shd)
    run = para.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    return para


def is_table_sep(line):
    s = line.strip().strip("|")
    return bool(s) and all(set(c.strip()) <= set("-: ") for c in s.split("|"))


def strip_md_inline(text):
    """Remove marcações inline de Markdown."""
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*",     r"\1", text)
    text = re.sub(r"__(.*?)__",     r"\1", text)
    text = re.sub(r"`(.*?)`",       r"\1", text)
    # Links [texto](url) → texto
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    return text


# ─────────────────────────────────────────────────────────────────────────────
# 5. PARSER MARKDOWN → DOCX
# ─────────────────────────────────────────────────────────────────────────────
def md_to_docx(doc, md_text):
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── Bloco de código ─────────────────────────────────────────────────
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                add_code_block(doc, "\n".join(code_lines))
            i += 1
            continue

        # ── Headings ────────────────────────────────────────────────────────
        m = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if m:
            level = min(len(m.group(1)), 4)
            title = strip_md_inline(m.group(2))
            add_heading(doc, title, level=level)
            i += 1
            continue

        # ── Linhas vazias ou apenas separadores (já removidos pelo clean_text)
        if not stripped:
            doc.add_paragraph("")
            i += 1
            continue

        # ── Tabela Markdown ─────────────────────────────────────────────────
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            if table_lines:
                header_cells = [c.strip() for c in table_lines[0].strip().strip("|").split("|")]
                body_rows = []
                for tl in table_lines[1:]:
                    if is_table_sep(tl):
                        continue
                    body_rows.append([c.strip() for c in tl.strip().strip("|").split("|")])
                n = len(header_cells)
                if n > 0:
                    tbl = doc.add_table(rows=1 + len(body_rows), cols=n)
                    tbl.style = "Table Grid"
                    hrow = tbl.rows[0]
                    for j, ct in enumerate(header_cells):
                        hrow.cells[j].text = strip_md_inline(ct)
                        for p in hrow.cells[j].paragraphs:
                            if p.runs:
                                p.runs[0].bold = True
                                p.runs[0].font.size = Pt(10)
                                p.runs[0].font.name = "Calibri"
                    for ri, row_data in enumerate(body_rows):
                        rw = tbl.rows[ri + 1]
                        for j, ct in enumerate(row_data[:n]):
                            rw.cells[j].text = strip_md_inline(ct)
                            for p in rw.cells[j].paragraphs:
                                for run in p.runs:
                                    run.font.size = Pt(10)
                                    run.font.name = "Calibri"
                doc.add_paragraph("")
            continue

        # ── Bullet list ─────────────────────────────────────────────────────
        if re.match(r"^[\-\*\+] ", stripped):
            text = re.sub(r"^[\-\*\+] ", "", stripped)
            text = strip_md_inline(text)
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(text)
            run.font.size = Pt(12)
            run.font.name = "Calibri"
            i += 1
            continue

        # ── Texto normal ────────────────────────────────────────────────────
        text = strip_md_inline(stripped)
        if text:
            add_body(doc, text)
        i += 1


# ─────────────────────────────────────────────────────────────────────────────
# 6. DOCUMENTO PRINCIPAL
# ─────────────────────────────────────────────────────────────────────────────
doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2)

# ══════════════════════════════════════════════════════════════════════════════
# 2.1 — MODELAGEM DO BANCO DE DADOS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2.1  Modelagem do Banco de Dados", level=1)

# --- 2.1.1 DER ---------------------------------------------------------------
add_heading(doc, "2.1.1  Modelo Conceitual — Diagrama Entidade-Relacionamento", level=2)
add_body(doc, "Abaixo o nosso Diagrama Entidade-Relacionamento:")
doc.add_picture(DER_PATH, width=Cm(14))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
p_leg = doc.add_paragraph(
    "Figura 1 — Diagrama Entidade-Relacionamento do Sistema de Quimioinformática Inteligente."
)
p_leg.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p_leg.runs:
    run.font.size = Pt(10)
    run.font.italic = True
doc.add_paragraph("")

# --- 2.1.2 Lógico ------------------------------------------------------------
add_heading(doc, "2.1.2  Modelo Lógico", level=2)
md_to_docx(doc, MODELO_LOGICO)
doc.add_paragraph("")

# --- 2.1.3 Físico ------------------------------------------------------------
add_heading(doc, "2.1.3  Modelo Físico (SQL)", level=2)
add_body(doc,
    "Script SQL completo para criação do banco de dados SQLite com todas as tabelas, "
    "índices, views e dados iniciais de teste:"
)
add_code_block(doc, MODELO_FISICO)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2.2 — DOCUMENTAÇÃO DA API
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2.2  Documentação da API (REST)", level=1)
md_to_docx(doc, API_DOCS)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2.3 — BACKEND IMPLEMENTADO
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2.3  Backend Implementado", level=1)
add_body(doc,
    "O código-fonte completo, estruturado no padrão MVC com FastAPI, foi versionado "
    "no GitHub e pode ser acessado no repositório oficial:"
)
p_link = doc.add_paragraph()
run_link = p_link.add_run("https://github.com/Alessandro-011/Quimiochat_IA")
run_link.font.color.rgb = RGBColor(0x00, 0x5a, 0xcc)
run_link.underline = True
run_link.font.size  = Pt(12)
run_link.font.name  = "Calibri"

doc.add_paragraph("")
add_body(doc,
    "A estrutura segue o padrão MVC com separação clara entre as camadas de modelos, "
    "rotas, controladores, schemas de validação e serviços. Abaixo a árvore de "
    "diretórios do backend:"
)
TREE = (
    "backend/\n"
    "  app/\n"
    "    controllers/   <- Lógica de negócio (CRUD, busca, autenticação)\n"
    "    database/      <- Configuração SQLAlchemy e sessão\n"
    "    middleware/    <- Verificação JWT (get_current_user)\n"
    "    models/        <- Entidades ORM (User, Molecule, Search...)\n"
    "    routes/        <- Roteadores FastAPI (auth, molecules, searches)\n"
    "    schemas/       <- Schemas Pydantic (validação entrada/saída)\n"
    "    services/      <- Serviços externos (Ollama, PubChem)\n"
    "    utils/         <- Utilitários (geração 3D, helpers)\n"
    "    main.py        <- Ponto de entrada da aplicação\n"
    "  requirements.txt\n"
    "  .env"
)
add_code_block(doc, TREE)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2.4 — INTEGRAÇÃO COM O FRONTEND
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2.4  Integração com o Frontend", level=1)
add_body(doc,
    "O frontend em Streamlit está consumindo as rotas do FastAPI em tempo real. "
    "A comunicação ocorre via requisições HTTP (biblioteca requests), com autenticação "
    "JWT armazenada na session_state do Streamlit para cada chamada à API."
)
doc.add_paragraph("")
p_note = doc.add_paragraph()
rn = p_note.add_run(
    "A interface possui tema Dark Mode personalizado e exibe os resultados comparativos "
    "da IA (Gemma2/Ollama) e do PubChem lado a lado, com medição de tempo de resposta "
    "e visualização 2D/3D interativa da molécula pesquisada."
)
rn.font.size    = Pt(11)
rn.font.name    = "Calibri"
rn.font.italic  = True
rn.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2.5 — README DO PROJETO
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2.5  Arquivo README do Projeto", level=1)
md_to_docx(doc, README)

# ─────────────────────────────────────────────────────────────────────────────
# SALVAR
# ─────────────────────────────────────────────────────────────────────────────
OUTPUT = os.path.join(DOCS_DIR, "Atividade2_Quimiochat.docx")
doc.save(OUTPUT)
print(f"Documento gerado com sucesso: {OUTPUT}")
